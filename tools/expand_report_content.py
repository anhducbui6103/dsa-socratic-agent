from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Duc_BA_20241582E_enhanced.docx"
ASSETS = ROOT / "docs" / "report-assets"


def paragraph_before(anchor, text: str, style) -> None:
    paragraph = anchor._parent.add_paragraph(text, style=style)
    anchor._p.addprevious(paragraph._p)


def image_before(doc: Document, anchor, image_name: str, caption: str, width: float = 5.9) -> None:
    image_path = ASSETS / image_name
    if not image_path.exists():
        return

    image_para = anchor._parent.add_paragraph(style=doc.styles["Normal"])
    image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_para.add_run().add_picture(str(image_path), width=Inches(width))
    anchor._p.addprevious(image_para._p)

    cap = anchor._parent.add_paragraph(caption, style=doc.styles["Normal"])
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    anchor._p.addprevious(cap._p)


def find_heading(doc: Document, text: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Heading not found: {text}")


def insert_block_before(doc: Document, heading_text: str, items: list[tuple[str, str | int]]) -> None:
    anchor = find_heading(doc, heading_text)
    styles = doc.styles
    for kind, value in items:
        if kind == "p":
            paragraph_before(anchor, str(value), styles["Normal"])
        elif kind == "h1":
            paragraph_before(anchor, str(value), styles["Heading 1"])
        elif kind == "h2":
            paragraph_before(anchor, str(value), styles["Heading 2"])
        elif kind == "h3":
            paragraph_before(anchor, str(value), styles["Heading 3"])
        elif kind == "img":
            name, caption = str(value).split("|", 1)
            image_before(doc, anchor, name, caption)


def main() -> None:
    doc = Document(DOCX)
    text = "\n".join(p.text for p in doc.paragraphs)
    if "Cơ sở lý thuyết" in text and "Minh họa node LangGraph trong code" in text:
        return

    insert_block_before(
        doc,
        "Mục tiêu",
        [
            ("h1", "Cơ sở lý thuyết"),
            ("h2", "Phương pháp Socratic trong dạy học thuật toán"),
            (
                "p",
                "Phương pháp Socratic là cách dạy học thông qua câu hỏi gợi mở thay vì truyền đạt trực tiếp đáp án. Trong bối cảnh học thuật toán, phương pháp này đặc biệt phù hợp vì mục tiêu không chỉ là biết một lời giải, mà là hiểu vì sao lời giải đó đúng, khi nào áp dụng được và giới hạn của nó nằm ở đâu.",
            ),
            (
                "p",
                "Một phản hồi Socratic tốt thường bắt đầu từ việc yêu cầu người học mô tả lại input, output, ràng buộc và ví dụ nhỏ. Sau đó hệ thống mới dẫn dắt tới việc xác định trạng thái cần lưu, thao tác cập nhật, invariant hoặc điều kiện dừng. Cách làm này giúp người học tự hình thành chuỗi suy luận thay vì sao chép lời giải.",
            ),
            (
                "p",
                "Trong hệ thống này, Socratic tutoring được hiện thực hóa bằng hint_level, policy không đưa full solution và Pedagogy Critic Agent. Hint level giúp mức độ hỗ trợ tăng dần; policy giữ phản hồi không quá trực tiếp; còn Pedagogy Critic kiểm tra phản hồi cuối trước khi gửi cho người học.",
            ),
            ("h2", "LLM Agent và workflow có trạng thái"),
            (
                "p",
                "LLM có khả năng sinh ngôn ngữ tự nhiên và hiểu ngữ cảnh tốt, nhưng nếu chỉ gọi LLM trực tiếp cho mỗi câu hỏi thì hệ thống khó kiểm soát hành vi, khó kiểm thử và dễ trả lời không nhất quán. Vì vậy, project sử dụng LLM như một thành phần trong workflow có trạng thái thay vì xem LLM là toàn bộ hệ thống.",
            ),
            (
                "p",
                "Một agent có trạng thái cần biết người học đang làm bài nào, đã xin bao nhiêu lần gợi ý, đã gửi code hay chưa, testcase nào từng được sinh và lần validation gần nhất cho kết quả gì. Các thông tin này được lưu trong LearningState để phản hồi ở lượt sau có tính liên tục.",
            ),
            ("h2", "Kiến trúc multi-agent"),
            (
                "p",
                "Kiến trúc multi-agent chia hệ thống thành nhiều thành phần chuyên trách. Thay vì một mô hình duy nhất vừa phân loại đề, vừa sinh testcase, vừa kiểm code, vừa đánh giá phản hồi, mỗi agent đảm nhiệm một vai trò rõ ràng. Cách tách này giúp giảm độ phức tạp trong prompt, dễ quan sát lỗi và dễ thay thế từng thành phần.",
            ),
            (
                "p",
                "Trong DSA Socratic Agent, AI Tutor Agent là agent điều phối trung tâm. Các agent còn lại gồm Intent Detector, Problem Classifier, Hint Generator, Testcase Generator, Code Execution Validator, Pedagogy Critic và Session Title Agent. Mỗi agent có input/output có cấu trúc để workflow có thể kết hợp kết quả một cách ổn định.",
            ),
            ("h2", "Vai trò của testcase trong học thuật toán"),
            (
                "p",
                "Testcase không chỉ dùng để chấm đúng sai. Trong quá trình học thuật toán, testcase là công cụ giúp người học kiểm tra giả định, phát hiện edge case và trace trạng thái qua từng bước. Một testcase nhỏ nhưng được chọn đúng có thể làm lộ lỗi tư duy rõ hơn nhiều so với việc đọc lời giải hoàn chỉnh.",
            ),
            (
                "p",
                "Vì vậy hệ thống phân biệt validation testcase và trace-only testcase. Validation testcase có expected output để chạy tự động. Trace-only testcase dùng khi expected output chưa đủ chắc chắn hoặc khi mục tiêu chính là hướng dẫn người học suy luận bằng tay.",
            ),
        ],
    )

    insert_block_before(
        doc,
        "Kiến trúc hệ thống",
        [
            ("h1", "Phương pháp thiết kế"),
            ("h2", "Nguyên tắc LLM-first nhưng có kiểm soát"),
            (
                "p",
                "Hệ thống ưu tiên dùng LLM cho các tác vụ cần hiểu ngôn ngữ tự nhiên như nhận diện intent, phân loại đề bài, sinh gợi ý và review phản hồi. Tuy nhiên, LLM không được gọi một cách tự do mà luôn nằm trong prompt policy, schema JSON và workflow có trạng thái.",
            ),
            (
                "p",
                "Các output quan trọng như classification, testsuite, validation result và pedagogy review được chuẩn hóa thành dataclass hoặc JSON schema. Việc chuẩn hóa này giúp hệ thống dễ debug, dễ viết test và dễ hiển thị thông tin trên UI.",
            ),
            ("h2", "Nguyên tắc không thay người học giải bài"),
            (
                "p",
                "Trong chế độ mặc định, hệ thống không cung cấp full code hoặc lời giải hoàn chỉnh. Nếu người học yêu cầu đáp án trực tiếp, workflow chuyển sang direct_solution_guard để biến yêu cầu đó thành gợi ý có kiểm soát. Đây là điểm khác biệt quan trọng giữa hệ thống tutor và chatbot giải bài.",
            ),
            (
                "p",
                "Mỗi phản hồi chính chỉ nên tập trung vào một hoặc hai điểm: một câu hỏi về ràng buộc, một edge case cần kiểm, một biến trạng thái cần trace hoặc một nhận xét về độ phức tạp. Cách này giữ cho quá trình học không bị quá tải thông tin.",
            ),
            ("h2", "Nguyên tắc quan sát được workflow"),
            (
                "p",
                "Agent Trace được thiết kế để hiển thị các bước kỹ thuật của workflow mà không tiết lộ chain-of-thought nội bộ của LLM. Trace giúp người phát triển biết node nào đã chạy, agent nào trả output gì và chất lượng phản hồi được kiểm tra ra sao.",
            ),
            (
                "p",
                "Trong bối cảnh vận hành thực tế, observability giúp phát hiện lỗi phân loại intent, lỗi testcase thiếu expected output, lỗi sandbox timeout hoặc phản hồi bị Pedagogy Critic đánh dấu rủi ro. Vì vậy Agent Trace vừa phục vụ demo, vừa là nền tảng debug hệ thống.",
            ),
        ],
    )

    insert_block_before(
        doc,
        "Learning Modules",
        [
            ("h2", "Minh họa node LangGraph trong code"),
            (
                "p",
                "Đoạn code dưới đây thể hiện cách workflow được tách thành các node độc lập. Mỗi node đại diện cho một hành vi chính của tutor agent, còn conditional edge đảm nhiệm việc định tuyến theo intent.",
            ),
            ("img", "code-langgraph-nodes.png|Hình 3: Trích đoạn khai báo các node LangGraph trong project"),
        ],
    )

    insert_block_before(
        doc,
        "Thiết kế giao diện",
        [
            ("h2", "Thiết kế dữ liệu trạng thái"),
            (
                "p",
                "LearningState là lớp dữ liệu trung tâm của hệ thống. Thay vì chỉ lưu lịch sử chat dạng text, state lưu cả thông tin học tập như dạng bài, concepts, hint level, attempts, testcase history, validation result và pedagogy flags. Đây là cơ sở để phản hồi của agent có tính cá nhân hóa.",
            ),
            (
                "p",
                "Nhờ LearningState, cùng một câu 'cho em thêm gợi ý' có thể tạo ra phản hồi khác nhau tùy vào bài toán hiện tại, số lần người học đã xin hint và kết quả validation gần nhất. Điều này giúp hệ thống khác với chatbot stateless thông thường.",
            ),
            ("img", "code-state-schema.png|Hình 6: Trích đoạn schema LearningState"),
        ],
    )

    insert_block_before(
        doc,
        "Luồng xử lý tiêu biểu",
        [
            ("h2", "Minh họa cơ chế review phản hồi"),
            (
                "p",
                "Trước khi phản hồi được gửi ra giao diện, Pedagogy Critic Agent kiểm tra mức độ an toàn sư phạm. Nếu phản hồi bị đánh dấu chưa an toàn, Tutor Agent yêu cầu LLM viết lại theo revision_instruction rồi review lại lần nữa.",
            ),
            (
                "p",
                "Cơ chế này giúp giảm rủi ro LLM trả lời quá trực tiếp, đưa full code hoặc bỏ qua hint level hiện tại. Đây là một ví dụ cho cách hệ thống dùng agent chất lượng để kiểm soát output của LLM.",
            ),
            ("img", "code-quality-review.png|Hình 8: Trích đoạn cơ chế Pedagogy Critic review và rewrite"),
        ],
    )

    insert_block_before(
        doc,
        "Đánh giá hệ thống",
        [
            ("h1", "Ví dụ minh họa kịch bản học tập"),
            ("h2", "Kịch bản 1: Người học gửi đề bài"),
            (
                "p",
                "Giả sử người học gửi bài toán Maximum Subarray: cho một mảng số nguyên, tìm tổng lớn nhất của một đoạn con liên tiếp. Hệ thống trước hết nhận diện đây là intent SUBMIT_PROBLEM, sau đó Problem Classifier phân loại bài toán vào nhóm array/dynamic programming với tín hiệu chính là đoạn con liên tiếp và tối ưu tổng.",
            ),
            (
                "p",
                "Ở bước tiếp theo, Testcase Generator Agent sinh một số case đại diện: mảng có cả số âm và số dương, mảng toàn số âm, mảng chỉ có một phần tử. Tutor Agent chưa đưa ngay thuật toán Kadane, mà hỏi người học cần theo dõi thông tin gì khi duyệt từng phần tử.",
            ),
            ("h2", "Kịch bản 2: Người học xin gợi ý nhiều lần"),
            (
                "p",
                "Khi người học xin gợi ý lần đầu, hint_level tăng từ 0 lên 1. Hệ thống chỉ gợi mở bằng câu hỏi về trạng thái cần lưu sau mỗi bước. Nếu người học tiếp tục xin gợi ý, hệ thống có thể nâng mức hỗ trợ bằng cách nhắc tới biến lưu tổng tốt nhất kết thúc tại vị trí hiện tại và biến lưu kết quả tốt nhất toàn cục.",
            ),
            (
                "p",
                "Dù hint_level tăng, Pedagogy Critic vẫn kiểm tra để phản hồi không biến thành lời giải hoàn chỉnh. Nhờ đó, hệ thống giữ được tinh thần tutor: hỗ trợ người học tiến gần lời giải nhưng vẫn để người học tự hoàn thiện.",
            ),
            ("h2", "Kịch bản 3: Người học gửi code sai edge case"),
            (
                "p",
                "Nếu người học gửi code khởi tạo tổng hiện tại bằng 0, thuật toán có thể sai với mảng toàn số âm. Code Execution Validator Agent hoặc testcase trace-only sẽ làm lộ vấn đề này. Tutor Agent không sửa code ngay, mà hỏi người học thử trace với input [-3, -2, -5] và quan sát biến khởi tạo ảnh hưởng thế nào tới kết quả.",
            ),
            (
                "p",
                "Cách phản hồi này giúp người học tự phát hiện lỗi thay vì chỉ nhận một bản sửa. Đây là mục tiêu chính của hệ thống: biến lỗi code thành cơ hội học về invariant, trạng thái và edge case.",
            ),
            ("h2", "Kịch bản 4: Người học yêu cầu lời giải trực tiếp"),
            (
                "p",
                "Khi người học yêu cầu full solution hoặc code mẫu, Intent Detector chuyển intent sang ASK_DIRECT_SOLUTION. Direct Solution Guard được kích hoạt để tránh đưa lời giải hoàn chỉnh. Thay vào đó, hệ thống đưa gợi ý cấp cao hơn hoặc đề xuất một testcase để người học tự kiểm.",
            ),
        ],
    )

    doc.save(DOCX)


if __name__ == "__main__":
    main()
