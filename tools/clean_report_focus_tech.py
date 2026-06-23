from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
INPUT_DOCX = ROOT / "Duc_BA_20241582E_enhanced.docx"
FALLBACK_DOCX = ROOT / "Duc_BA_20241582E_tech_report.docx"
ASSETS = ROOT / "docs" / "report-assets"


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def add_para(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style=doc.styles["Normal"])


def add_heading(doc: Document, text: str, level: int) -> None:
    doc.add_paragraph(text, style=doc.styles[f"Heading {level}"])


def add_image(doc: Document, name: str, caption: str, width: float = 5.9) -> None:
    path = ASSETS / name
    if not path.exists():
        return
    p = doc.add_paragraph(style=doc.styles["Normal"])
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption, style=doc.styles["Normal"])
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def main() -> None:
    doc = Document(INPUT_DOCX)

    start = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip().startswith("TÓM TẮT"))
    for paragraph in list(doc.paragraphs[start:]):
        delete_paragraph(paragraph)

    add_para(
        doc,
        "TÓM TẮT— Báo cáo trình bày hệ thống DSA Socratic Agent, một LLM-powered multi-agent tutor hỗ trợ học Cấu trúc dữ liệu và Giải thuật. Hệ thống kết hợp mô hình ngôn ngữ lớn, LangGraph, Streamlit và các agent kiểm soát chất lượng để dẫn dắt người học theo phương pháp Socratic. Thay vì cung cấp lời giải trực tiếp, hệ thống phân tích ngữ cảnh, duy trì trạng thái học tập, sinh gợi ý theo mức độ, tạo testcase, kiểm tra code và đánh giá phản hồi trước khi gửi tới người học.",
    )

    add_heading(doc, "Giới thiệu", 1)
    add_para(
        doc,
        "Cấu trúc dữ liệu và Giải thuật là nhóm kiến thức nền tảng trong khoa học máy tính, nhưng quá trình học thường đòi hỏi nhiều hơn việc ghi nhớ thuật toán. Người học cần biết phân tích đề bài, xác định ràng buộc, lựa chọn cấu trúc dữ liệu phù hợp, kiểm tra edge case và đánh giá độ phức tạp. Đây là những kỹ năng khó hình thành nếu chỉ nhận lời giải hoàn chỉnh.",
    )
    add_para(
        doc,
        "DSA Socratic Agent được thiết kế như một trợ giảng AI định hướng tư duy. Hệ thống sử dụng LLM để hiểu ngữ cảnh và sinh phản hồi tự nhiên, đồng thời sử dụng workflow agentic để chia quá trình xử lý thành các bước có cấu trúc. Cách tiếp cận này giúp kết hợp sự linh hoạt của LLM với khả năng kiểm soát của các agent chuyên trách.",
    )
    add_para(
        doc,
        "Trọng tâm của hệ thống là kiến trúc multi-agent: AI Tutor Agent giữ vai trò điều phối, các learning modules xử lý nhiệm vụ chuyên môn, còn quality agents kiểm tra testcase, kết quả chạy code và mức độ phù hợp sư phạm của phản hồi.",
    )

    add_heading(doc, "Mục tiêu", 1)
    add_heading(doc, "Mục tiêu tổng quát", 2)
    add_para(
        doc,
        "Xây dựng một hệ thống AI Tutor hỗ trợ học Cấu trúc dữ liệu và Giải thuật bằng phương pháp Socratic, có khả năng hiểu hội thoại, duy trì trạng thái học tập, điều phối nhiều agent chuyên trách và kiểm soát chất lượng phản hồi.",
    )
    add_heading(doc, "Mục tiêu cụ thể", 2)
    add_heading(doc, "Điều phối hội thoại bằng LangGraph", 3)
    add_para(
        doc,
        "Workflow của hệ thống được mô hình hóa bằng LangGraph với các node độc lập theo intent. Việc tách node giúp pipeline rõ ràng, dễ mở rộng, dễ kiểm thử và dễ quan sát trong quá trình vận hành.",
    )
    add_heading(doc, "Ứng dụng LLM trong các tác vụ ngữ cảnh", 3)
    add_para(
        doc,
        "LLM được sử dụng cho intent detection, problem classification, hint generation, testcase generation, pedagogy review và final response generation. Các guideline nội bộ chỉ đóng vai trò bổ trợ ngữ cảnh cho prompt, không thay thế khả năng sinh phản hồi của LLM.",
    )
    add_heading(doc, "Xây dựng Quality Layer", 3)
    add_para(
        doc,
        "Quality Layer được thiết kế để tăng độ tin cậy của hệ thống. Testcase Generator Agent tạo case kiểm thử và case trace; Code Execution Validator Agent kiểm tra code khi có expected output; Pedagogy Critic Agent đánh giá phản hồi theo chính sách Socratic.",
    )
    add_heading(doc, "Cung cấp giao diện tương tác", 3)
    add_para(
        doc,
        "Giao diện Streamlit cung cấp môi trường chat trực quan, hỗ trợ nhiều phiên học, hiển thị trạng thái học tập và Agent Trace. Giao diện đóng vai trò lớp trình diễn và tương tác, trong khi logic chính nằm ở core agent.",
    )

    add_heading(doc, "Công nghệ sử dụng", 1)
    add_para(
        doc,
        "Hệ thống sử dụng Python cho phần lõi xử lý agent và định nghĩa dữ liệu. Python phù hợp với bài toán này nhờ hệ sinh thái mạnh cho AI, xử lý văn bản, kiểm thử và xây dựng prototype nhanh.",
    )
    add_para(
        doc,
        "LangGraph được sử dụng để tổ chức workflow agentic. Thay vì xử lý toàn bộ hội thoại bằng một hàm duy nhất, LangGraph biểu diễn quá trình xử lý thành các node và cạnh điều kiện, giúp hệ thống có cấu trúc rõ ràng hơn.",
    )
    add_para(
        doc,
        "Gemini API được sử dụng như LLM provider để sinh phản hồi, phân loại intent, phân loại bài toán, tạo testcase và review chất lượng sư phạm. Việc tách lớp LLM client giúp hệ thống có thể thay đổi provider trong tương lai mà không ảnh hưởng nhiều tới orchestration.",
    )
    add_para(
        doc,
        "Streamlit được sử dụng để xây dựng giao diện chat. Công nghệ này cho phép tạo UI tương tác nhanh, phù hợp với hệ thống AI prototype có nhiều trạng thái và cần hiển thị kết quả xử lý theo thời gian thực.",
    )
    add_para(
        doc,
        "Các thành phần kiểm thử sử dụng unittest nhằm đảm bảo các schema, validator và logic điều phối cốt lõi hoạt động ổn định. Sandbox execution dùng subprocess có giới hạn timeout để chạy code Python trên testcase khi đủ dữ liệu đầu vào và expected output.",
    )

    add_heading(doc, "Kiến trúc hệ thống", 1)
    add_heading(doc, "Tổng quan kiến trúc", 2)
    add_para(
        doc,
        "Hệ thống gồm năm lớp chính: User Interface, Tutor Orchestration Layer, Learning Modules, Quality Layer và State & Memory Layer. Các lớp này tách biệt trách nhiệm giữa giao diện, điều phối, xử lý chuyên môn, kiểm soát chất lượng và lưu trạng thái học tập.",
    )
    add_image(doc, "architecture.png", "Hình 1: Tổng quan kiến trúc hệ thống multi-agent", 5.9)

    add_heading(doc, "Tutor Orchestration Layer", 2)
    add_para(
        doc,
        "AI Tutor Agent là thành phần điều phối trung tâm. Agent nhận message từ người học, gọi LLM để xác định intent, cập nhật LearningState, định tuyến tới module hoặc agent phù hợp, sau đó tổng hợp kết quả thành phản hồi cuối.",
    )
    add_para(
        doc,
        "Workflow LangGraph gồm các node detect_intent, submit_problem, request_hint, submit_code, direct_solution_guard, submit_approach, ask_theory và finalize. Mỗi node xử lý một loại hành động cụ thể và trả output có cấu trúc để các bước sau sử dụng.",
    )
    add_image(doc, "langgraph.png", "Hình 2: Workflow LangGraph tách node theo intent", 5.9)

    add_heading(doc, "Learning Modules", 2)
    add_para(
        doc,
        "Learning Modules cung cấp năng lực chuyên môn cho quá trình dạy học. Problem Classifier nhận diện chủ đề và pattern thuật toán; Hint Generator sinh gợi ý theo hint level; Code Analyzer phân tích logic, độ phức tạp và edge case; Knowledge Context cung cấp taxonomy và rubric nội bộ cho prompt.",
    )
    add_para(
        doc,
        "Các module này không hoạt động độc lập như các chatbot nhỏ riêng lẻ, mà được Tutor Agent gọi theo ngữ cảnh. Kết quả của module được lưu vào state hoặc truyền sang Quality Layer để kiểm tra thêm.",
    )

    add_heading(doc, "Quality Layer", 2)
    add_para(
        doc,
        "Quality Layer gồm ba agent chính. Testcase Generator Agent sinh TestSuite với các case basic, edge, adversarial hoặc stress. Code Execution Validator Agent chạy code Python trên testcase có expected output. Pedagogy Critic Agent đánh giá phản hồi nháp trước khi gửi cho người học.",
    )
    add_para(
        doc,
        "Testcase được chia thành hai nhóm: validation testcase và trace-only testcase. Validation testcase có expected output và dùng để chạy kiểm tra tự động. Trace-only testcase dùng để hướng dẫn người học tự kiểm tra logic khi đề bài chưa đủ rõ để tạo oracle tin cậy.",
    )
    add_para(
        doc,
        "Pedagogy Critic kiểm tra các rủi ro như lộ full solution, phản hồi quá trực tiếp, không phù hợp hint level hoặc đưa quá nhiều gợi ý trong một lượt. Nếu phản hồi chưa đạt, hệ thống yêu cầu LLM viết lại trước khi gửi.",
    )

    add_heading(doc, "State & Memory Layer", 2)
    add_para(
        doc,
        "LearningState lưu trữ session_id, current_problem, problem_type, concepts, hint_level, student_attempts, misconceptions, generated_tests, latest_validation, pedagogy_flags và next_action. Nhờ state, hệ thống có thể cá nhân hóa phản hồi theo tiến trình học thay vì xử lý từng câu hỏi độc lập.",
    )
    add_para(
        doc,
        "Session Title Agent sinh tiêu đề ngắn cho từng phiên chat dựa trên nội dung trao đổi ban đầu. Cách này giúp người dùng quản lý nhiều phiên học hiệu quả hơn và tránh hiển thị nguyên văn message dài ở danh sách phiên.",
    )

    add_heading(doc, "Agent Trace và Observability", 2)
    add_para(
        doc,
        "Agent Trace ghi lại các bước workflow đã chạy trong một lượt hội thoại, bao gồm node name, trạng thái, summary và output rút gọn. Trace này phục vụ debug, demo và đánh giá hệ thống; nó không phải là chain-of-thought nội bộ của LLM.",
    )
    add_image(doc, "agent-trace.png", "Hình 3: Minh họa Agent Trace trong một lượt hội thoại", 5.9)

    add_heading(doc, "Thiết kế giao diện", 1)
    add_para(
        doc,
        "Giao diện được thiết kế theo mô hình chat. Cột trái hiển thị danh sách phiên học và tiêu đề ngắn của từng phiên. Vùng trung tâm hiển thị hội thoại hiện tại. Ô nhập được đặt cố định phía dưới để người dùng có thể nhập câu hỏi, đề bài hoặc code một cách liên tục.",
    )
    add_para(
        doc,
        "Panel trạng thái bên phải hiển thị các thông tin phục vụ học tập và quan sát hệ thống như hint level, problem type, concepts, testcase summary, validation result, pedagogy review và Agent Trace. Những chi tiết cấu hình hạ tầng không được đưa vào UI để giữ giao diện tập trung vào trải nghiệm học.",
    )
    add_image(doc, "demo-ui.png", "Hình 4: Minh họa giao diện Streamlit", 5.9)

    add_heading(doc, "Thiết kế các agent", 1)
    add_heading(doc, "Intent Detector Agent", 2)
    add_para(
        doc,
        "Intent Detector Agent phân loại message thành các intent như ASK_THEORY, SUBMIT_PROBLEM, REQUEST_HINT, SUBMIT_APPROACH, SUBMIT_CODE và ASK_DIRECT_SOLUTION. Intent là tín hiệu đầu vào quan trọng để LangGraph chọn nhánh xử lý.",
    )
    add_heading(doc, "Problem Classifier Agent", 2)
    add_para(
        doc,
        "Problem Classifier Agent phân tích đề bài và trả về topic, pattern, confidence, key_signals và recommended_hint_path. Output dạng JSON giúp các module tiếp theo sử dụng kết quả một cách nhất quán.",
    )
    add_heading(doc, "Hint Generator Agent", 2)
    add_para(
        doc,
        "Hint Generator Agent sinh gợi ý dựa trên hint_level, current_problem, attempts gần nhất và policy Socratic. Ở hint level thấp, agent ưu tiên câu hỏi định hướng; ở level cao hơn, agent có thể gợi ý cấu trúc dữ liệu, invariant hoặc edge case cần kiểm tra.",
    )
    add_heading(doc, "Testcase Generator Agent", 2)
    add_para(
        doc,
        "Testcase Generator Agent tạo TestSuite gồm nhiều GeneratedTestCase. Mỗi testcase có tên, input, expected output nếu xác định được, category và rationale. Testcase vừa phục vụ validation, vừa hỗ trợ người học trace logic.",
    )
    add_heading(doc, "Code Execution Validator Agent", 2)
    add_para(
        doc,
        "Code Execution Validator Agent chạy code Python trên testcase có expected output và chuẩn hóa kết quả thành PASSED, FAILED, RUNTIME_ERROR, TIMEOUT hoặc SKIPPED. Kết quả validation được Tutor Agent diễn giải thành gợi ý học tập thay vì trả log thô.",
    )
    add_heading(doc, "Pedagogy Critic Agent", 2)
    add_para(
        doc,
        "Pedagogy Critic Agent kiểm tra phản hồi nháp theo tiêu chí sư phạm. Agent này giúp hệ thống duy trì phong cách Socratic, tránh cung cấp lời giải quá sớm và đảm bảo mỗi lượt phản hồi tập trung vào một số gợi ý chính.",
    )
    add_heading(doc, "Session Title Agent", 2)
    add_para(
        doc,
        "Session Title Agent sinh tiêu đề ngắn cho phiên học, ví dụ Maximum Subarray, DFS trên đồ thị hoặc Sliding Window. Đây là agent hỗ trợ trải nghiệm người dùng và quản lý nhiều phiên hội thoại.",
    )

    add_heading(doc, "Luồng xử lý tiêu biểu", 1)
    add_para(
        doc,
        "Khi người học gửi đề bài, hệ thống detect intent, phân loại bài toán, sinh testcase gợi ý và hỏi người học làm rõ input/output hoặc ràng buộc. Khi người học xin gợi ý, hint_level tăng và Hint Generator sinh phản hồi mới. Khi người học gửi code, Code Analyzer và Execution Validator cung cấp tín hiệu để Tutor Agent đặt câu hỏi dẫn dắt.",
    )
    add_para(
        doc,
        "Nếu người học yêu cầu lời giải trực tiếp, direct_solution_guard được kích hoạt. Hệ thống không trả full code trong chế độ mặc định mà chuyển yêu cầu thành gợi ý cấp cao hơn, testcase tự kiểm hoặc câu hỏi giúp người học tự hoàn thiện lời giải.",
    )
    add_image(doc, "code-pipeline.png", "Hình 5: Trích đoạn pipeline LLM-first trong Tutor Orchestrator", 5.9)

    add_heading(doc, "Đánh giá hệ thống", 1)
    add_para(
        doc,
        "Hệ thống được đánh giá qua các kịch bản hội thoại phổ biến trong học DSA: gửi bài Maximum Subarray, xin hint nhiều lần, gửi code sai edge case toàn số âm, gửi câu hỏi lý thuyết và yêu cầu full solution. Các kịch bản này kiểm tra khả năng điều phối agent, duy trì state và kiểm soát phản hồi.",
    )
    add_para(
        doc,
        "Tiêu chí đánh giá gồm: phản hồi không lộ full solution, hint phù hợp với mức độ hiện tại, testcase có rationale rõ ràng, validation result được diễn giải thành gợi ý học tập và Agent Trace thể hiện đúng các bước workflow đã chạy.",
    )
    add_para(
        doc,
        "Bộ kiểm thử tự động kiểm tra các thành phần cốt lõi như intent handling, testcase schema, pedagogy review và execution validator. Kiểm thử giúp đảm bảo các output có cấu trúc ổn định để giao diện và workflow có thể sử dụng.",
    )

    add_heading(doc, "Hướng phát triển", 1)
    add_para(
        doc,
        "Hướng phát triển tiếp theo gồm lưu trữ session bằng database, mở rộng sandbox bằng môi trường cách ly mạnh hơn, bổ sung dashboard đánh giá chất lượng, hỗ trợ thêm ngôn ngữ lập trình và xây dựng benchmark bài tập DSA để đo chất lượng phản hồi.",
    )
    add_para(
        doc,
        "Ngoài ra, hệ thống có thể bổ sung Evaluation Agent để tự chấm điểm phản hồi theo rubric, tích hợp tracing bền vững cho từng phiên học và cải thiện cơ chế sinh expected output bằng oracle nội bộ hoặc reference solution không hiển thị cho người học.",
    )

    add_heading(doc, "Kết luận", 1)
    add_para(
        doc,
        "DSA Socratic Agent là một hệ thống multi-agent tutor ứng dụng LLM vào hỗ trợ học Cấu trúc dữ liệu và Giải thuật. Kiến trúc kết hợp AI Tutor Agent, LangGraph workflow, Learning Modules, Quality Layer và State Management để tạo ra phản hồi có ngữ cảnh, có kiểm soát và phù hợp với mục tiêu học tập.",
    )
    add_para(
        doc,
        "So với chatbot hỏi đáp thông thường, hệ thống có điểm mạnh ở khả năng điều phối nhiều agent, lưu trạng thái học tập, sinh testcase, kiểm tra code và đánh giá phản hồi trước khi gửi. Đây là nền tảng phù hợp để tiếp tục phát triển thành một trợ giảng AI có khả năng cá nhân hóa quá trình học DSA.",
    )

    try:
        doc.save(INPUT_DOCX)
    except PermissionError:
        doc.save(FALLBACK_DOCX)


if __name__ == "__main__":
    main()
