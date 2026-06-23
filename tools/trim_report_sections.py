from __future__ import annotations

from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Duc_BA_20241582E_enhanced.docx"
OUTPUT_DOCX = ROOT / "Duc_BA_20241582E_final.docx"

REMOVE_H1 = {
    "Phương pháp thiết kế",
    "Thiết kế giao diện",
    "Thiết kế các agent",
    "Luồng xử lý tiêu biểu",
    "Ví dụ minh họa kịch bản học tập",
    "Đánh giá hệ thống",
}


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def is_heading_1(paragraph) -> bool:
    return paragraph.style.name == "Heading 1"


def section_range(doc: Document, heading_text: str) -> tuple[int, int]:
    start = next(i for i, p in enumerate(doc.paragraphs) if is_heading_1(p) and p.text.strip() == heading_text)
    end = len(doc.paragraphs)
    for i in range(start + 1, len(doc.paragraphs)):
        if is_heading_1(doc.paragraphs[i]):
            end = i
            break
    return start, end


def collect_section_body(doc: Document, heading_text: str) -> list[tuple[str, str]]:
    start, end = section_range(doc, heading_text)
    return [(p.text, p.style.name) for p in doc.paragraphs[start + 1 : end] if p.text.strip()]


def remove_section(doc: Document, heading_text: str) -> None:
    try:
        start, end = section_range(doc, heading_text)
    except StopIteration:
        return
    for paragraph in list(doc.paragraphs[start:end]):
        delete_paragraph(paragraph)


def insert_after_section_heading(doc: Document, heading_text: str, items: list[tuple[str, str]]) -> None:
    heading = next(p for p in doc.paragraphs if is_heading_1(p) and p.text.strip() == heading_text)
    anchor = heading

    # Insert in reverse because each new paragraph is placed immediately after
    # the heading anchor.
    for text, style_name in reversed(items):
        paragraph = doc.add_paragraph(text, style=doc.styles[style_name])
        anchor._p.addnext(paragraph._p)


def main() -> None:
    doc = Document(DOCX)

    development_items = collect_section_body(doc, "Hướng phát triển")
    remove_section(doc, "Hướng phát triển")

    for heading in REMOVE_H1:
        remove_section(doc, heading)

    if development_items:
        merged_items = [("Hướng phát triển tiếp theo", "Heading 2"), *development_items]
        insert_after_section_heading(doc, "Kết luận", merged_items)

    try:
        doc.save(DOCX)
    except PermissionError:
        doc.save(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
