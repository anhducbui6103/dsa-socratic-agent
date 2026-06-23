from __future__ import annotations

from copy import deepcopy
from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "Duc_BA_20241582E.original.docx"
TARGET = ROOT / "Duc_BA_20241582E.docx"
ASSETS = ROOT / "docs" / "report-assets"


def set_text(paragraph, text: str) -> None:
    """Replace paragraph text while preserving the paragraph style and first-run formatting."""
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def insert_paragraph_after(paragraph, text: str = "", style=None):
    new_p = deepcopy(paragraph._p)
    paragraph._p.addnext(new_p)
    inserted = paragraph._parent.paragraphs[
        [p._p for p in paragraph._parent.paragraphs].index(new_p)
    ]
    for run in inserted.runs:
        run.text = ""
    if style is not None:
        inserted.style = style
    if text:
        if inserted.runs:
            inserted.runs[0].text = text
        else:
            inserted.add_run(text)
    return inserted


def add_image_to_paragraph(paragraph, image_path: Path, width: float = 5.9) -> None:
    for run in paragraph.runs:
        run.text = ""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(width))


def main() -> None:
    template = SOURCE if SOURCE.exists() else TARGET
    if template != TARGET:
        shutil.copy2(template, TARGET)
    doc = Document(TARGET)
    p = doc.paragraphs

    replacements = {
        5: "ĐỒ ÁN 2",
        7: "AI Agent hỗ trợ học Cấu trúc dữ liệu và Giải thuật\n",
        9: "Hà Nội, Ngày 22 tháng 6 năm 2026",
        11: "Mục lục",
        13: (
            "TÓM TẮT— Đề tài xây dựng một hệ thống LLM-powered Multi-Agent Tutor hỗ trợ sinh viên học môn "
            "Cấu trúc dữ liệu và Giải thuật theo phương pháp Socratic. Hệ thống không tập trung cung cấp lời giải trực tiếp, "
            "mà dẫn dắt người học thông qua câu hỏi gợi mở, gợi ý nhiều cấp độ, testcase tự kiểm và phản hồi dựa trên trạng thái học tập. "
            "Prototype hiện tại được triển khai bằng Python, Streamlit, LangGraph và Gemini API. Bên cạnh AI Tutor Agent đóng vai trò điều phối, "
            "hệ thống bổ sung Quality Layer gồm Testcase Generator Agent, Code Execution Validator Agent và Pedagogy Critic Agent để kiểm soát chất lượng "
            "trước khi phản hồi cho sinh viên. Kết quả là một prototype chạy được, có giao diện chat nhiều phiên, sandbox Python bước đầu và bộ kiểm thử tự động."
        ),
        18: (
            "Xuất phát từ thực tế đó, đề tài xây dựng DSA Socratic Agent - một hệ thống AI Agent hỗ trợ học Cấu trúc dữ liệu và Giải thuật "
            "theo hướng phát triển tư duy. Hệ thống sử dụng LLM để hiểu ngữ cảnh, nhận diện intent, phân loại bài toán, sinh testcase và tạo phản hồi. "
            "Tuy nhiên, LLM không được dùng một cách tự do mà nằm trong pipeline có quản lý trạng thái, policy Socratic và kiểm tra sư phạm. "
            "Nhờ vậy, hệ thống hướng tới vai trò trợ giảng thông minh: hỗ trợ người học suy luận từng bước, nhưng không làm thay toàn bộ bài toán."
        ),
        21: (
            "Xây dựng một prototype AI Agent hỗ trợ sinh viên học tập môn Cấu trúc dữ liệu và Giải thuật theo hướng phát triển tư duy, "
            "thông qua việc dùng LLM và multi-agent workflow để hướng dẫn bằng phương pháp Socratic thay vì cung cấp lời giải trực tiếp."
        ),
        23: "Xây dựng cơ chế nhận diện intent và phân loại bài toán bằng LLM",
        24: (
            "Hệ thống sử dụng Gemini API để nhận diện intent của người học, ví dụ: gửi đề bài, xin gợi ý, gửi code, trình bày ý tưởng hoặc yêu cầu lời giải trực tiếp. "
            "Sau đó LLM Problem Classifier phân tích đề bài để dự đoán nhóm thuật toán như sliding window, graph, greedy, dynamic programming hoặc sorting/searching. "
            "Kết quả phân loại được trả về dưới dạng có cấu trúc gồm topic, pattern, confidence, key signals và recommended hint path."
        ),
        26: (
            "Hệ thống xây dựng cơ chế gợi ý tăng dần dựa trên `hint_level`. Khi sinh viên yêu cầu gợi ý, agent tăng mức gợi ý và yêu cầu LLM sinh phản hồi dựa trên trạng thái hiện tại. "
            "Các gợi ý vẫn tuân thủ nguyên tắc Socratic: ưu tiên câu hỏi định hướng, ví dụ nhỏ, edge case hoặc yêu cầu trace thay vì đưa ngay công thức hoặc code hoàn chỉnh."
        ),
        28: (
            "AI Tutor Agent đóng vai trò điều phối hội thoại. Agent nhận input, gọi LLM để xác định intent, cập nhật LearningState, lựa chọn module hoặc quality agent phù hợp, "
            "sau đó sinh phản hồi cuối. Agent duy trì trạng thái gồm bài toán hiện tại, dạng bài, hint level, lịch sử attempts, testcase đã sinh, kết quả validation và pedagogy review."
        ),
        30: (
            "Hệ thống tiếp nhận code do sinh viên gửi và phân tích ở mức tư duy giải bài, độ phức tạp và edge cases. Với Python, hệ thống có Code Execution Validator Agent để chạy code trên testcase khi có expected output. "
            "Kết quả chạy code không được dùng để đưa đáp án trực tiếp mà được chuyển thành câu hỏi hoặc gợi ý giúp sinh viên tự phát hiện lỗi."
        ),
        31: "Bổ sung Quality Layer cho hệ thống multi-agent",
        32: (
            "Trong giai đoạn prototype, hệ thống tập trung vào Quality Layer gồm ba agent: Testcase Generator Agent, Code Execution Validator Agent và Pedagogy Critic Agent. "
            "Testcase Generator sinh case basic, edge, adversarial hoặc stress để hỗ trợ trace; Execution Validator chạy code Python có timeout; Pedagogy Critic kiểm phản hồi có lộ lời giải, quá trực tiếp hoặc sai mức gợi ý hay không. "
            "Thiết kế này phù hợp với mục tiêu hiện tại là kiểm soát chất lượng phản hồi và chứng minh luồng multi-agent."
        ),
        35: (
            "Hệ thống được thiết kế theo kiến trúc LLM-powered multi-agent. Một AI Tutor Agent trung tâm điều phối toàn bộ quá trình, trong khi các agent chuyên trách đảm nhiệm từng nhiệm vụ cụ thể như phân loại bài toán, sinh testcase, validate code và kiểm tra sư phạm. "
            "Cách tổ chức này giúp hệ thống không chỉ là chatbot trả lời một lượt, mà có thể theo dõi tiến trình học và phản hồi theo ngữ cảnh."
        ),
        36: (
            "Về tổng thể, hệ thống gồm năm lớp: giao diện người dùng bằng Streamlit, lớp điều phối Tutor Orchestration, các learning modules, Quality Layer và State & Memory Layer. "
            "Khi người học gửi input, Tutor Agent gọi LLM để nhận diện intent, cập nhật trạng thái, gọi module phù hợp, sinh phản hồi nháp, đưa qua Pedagogy Critic và chỉ gửi phản hồi cuối sau khi đáp ứng policy Socratic."
        ),
        38: "Hình 1: Tổng quan kiến trúc hệ thống multi-agent hiện tại",
        40: (
            "AI Tutor Agent là thành phần trung tâm, chịu trách nhiệm điều phối và ra quyết định trong toàn bộ quá trình tương tác. "
            "Trong code, thành phần này được hiện thực bởi lớp `DsaLearningAgent` trong file `src/dsa_agent/orchestrator.py`. "
            "Agent sử dụng LLM cho các bước cần hiểu ngữ cảnh và sử dụng module Python cho các bước cần tính quyết định như validation hoặc quản lý state."
        ),
        41: (
            "Pipeline xử lý chính gồm: LLM detect intent, cập nhật LearningState, gọi LLM classifier hoặc testcase generator khi cần, chạy validator nếu sinh viên gửi code, sinh phản hồi bằng LLM, review bằng Pedagogy Critic và rewrite nếu phản hồi chưa an toàn. "
            "Nguyên tắc quan trọng nhất là không cung cấp full code hoặc lời giải hoàn chỉnh trong chế độ mặc định."
        ),
        43: (
            "Orchestrator Agent hiện thực hóa cơ chế ra quyết định của AI Tutor Agent. Dựa trên intent, orchestrator lựa chọn nhánh xử lý: gửi đề bài, xin gợi ý, gửi code, trình bày ý tưởng, hỏi lý thuyết hoặc yêu cầu lời giải trực tiếp."
        ),
        44: (
            "Cụ thể, khi sinh viên gửi đề bài, orchestrator gọi LLM Problem Classifier và Testcase Generator. Khi sinh viên xin gợi ý, hệ thống tăng hint level và sinh hint mới. "
            "Khi sinh viên gửi code, orchestrator gọi Execution Validator và tạo phản hồi theo hướng trace hoặc edge case. Khi người học đòi lời giải, agent chuyển yêu cầu thành gợi ý Socratic thay vì đưa đáp án."
        ),
        46: (
            "State Management lưu trữ trạng thái học tập trong từng phiên chat. Trạng thái bao gồm `session_id`, `current_problem`, `problem_type`, `concepts`, `hint_level`, `student_attempts`, `generated_tests`, `latest_validation`, `pedagogy_flags` và `next_action`."
        ),
        47: (
            "Sau mỗi lượt hội thoại, state được cập nhật để agent hiểu người học đang ở bước nào. Ví dụ, `hint_level` tăng khi sinh viên xin gợi ý, `generated_tests` được cập nhật khi hệ thống sinh testcase, và `latest_validation` lưu kết quả chạy code gần nhất."
        ),
        48: "Thiết kế các module và quality agents",
        50: (
            "Problem Classifier hiện dùng LLM để phân loại bài toán. Prompt yêu cầu LLM trả JSON hợp lệ với topic, pattern, confidence, key signals và recommended hint path. "
            "Cách này linh hoạt hơn keyword matching vì có thể hiểu mô tả đề bài tự nhiên hơn."
        ),
        51: "Testcase Generator Agent",
        52: (
            "Testcase Generator Agent dùng LLM để sinh `TestSuite` gồm nhiều `GeneratedTestCase`. Mỗi testcase có name, input, expected output, category và rationale. "
            "Trong prototype, nhiều testcase có `expected_output = null` vì mục tiêu chính là hỗ trợ sinh viên trace tay; ở các phiên bản sau có thể bổ sung oracle để chấm tự động tốt hơn."
        ),
        53: "Hint Generation Module",
        54: (
            "Hint Generation Module không còn dùng câu gợi ý cố định làm phản hồi chính. Thay vào đó, LLM sinh gợi ý dựa trên LearningState, intent, hint level và policy Socratic. "
            "Các guideline trong code chỉ đóng vai trò định hướng prompt."
        ),
        55: "Code Analyzer và Execution Validator",
        56: (
            "Code Analyzer cung cấp nhận xét ở mức logic, độ phức tạp và edge case. Execution Validator chạy Python code bằng subprocess với timeout theo từng test, trả về các trạng thái như PASSED, FAILED, RUNTIME_ERROR, TIMEOUT hoặc SKIPPED. "
            "Kết quả này giúp tutor đặt câu hỏi chính xác hơn thay vì sửa code trực tiếp."
        ),
        59: (
            "Prototype hiện tại được triển khai bằng Python. Giao diện demo sử dụng Streamlit, workflow agentic được bọc bằng LangGraph, LLM provider là Gemini API và cấu hình được đọc từ file `.env` thông qua python-dotenv."
        ),
        60: (
            "Trong phạm vi Đồ án 2, hệ thống không xây dựng backend riêng như Node.js, NestJS hoặc FastAPI. Streamlit đảm nhiệm vai trò giao diện demo và quản lý session ở mức memory. "
            "Điều này giúp tập trung vào phần lõi agentic AI và phù hợp với quy mô prototype."
        ),
        61: (
            "Mô hình ngôn ngữ lớn được sử dụng thông qua Gemini API. LLM đảm nhiệm các tác vụ cần hiểu ngữ cảnh như intent detection, problem classification, testcase generation, final response và pedagogy review."
        ),
        62: (
            "Hệ thống hiện không triển khai module truy xuất tài liệu. Phần knowledge context trong prototype chủ yếu đến từ prompt, taxonomy và các tài liệu thiết kế nội bộ. "
            "Nhờ đó phạm vi đồ án tập trung hơn vào agent orchestration, quality layer và trải nghiệm demo."
        ),
        63: (
            "Các thư viện chính gồm LangGraph cho workflow orchestration, Streamlit cho UI, python-dotenv cho cấu hình môi trường và unittest cho kiểm thử. "
            "API key được đặt trong `.env` dưới biến `GEMINI_API_KEY` và không nhập trực tiếp trên giao diện."
        ),
        65: (
            "AI Agent được triển khai trong lớp `DsaLearningAgent`. Khi nhận input, agent gọi LLM để phân loại intent, ghi nhận input vào `student_attempts`, sau đó chọn nhánh xử lý phù hợp."
        ),
        66: (
            "Cơ chế hoạt động của agent có thể xem như vòng lặp observe - reason - act: quan sát input và state, suy luận nhánh xử lý, gọi module/agent cần thiết và tạo phản hồi theo policy Socratic."
        ),
        67: (
            "Ở bước hành động, agent có thể gọi LLM classifier, LLM testcase generator, execution validator hoặc pedagogy critic. "
            "Phản hồi cuối luôn ưu tiên dẫn dắt tư duy, hạn chế cung cấp lời giải trực tiếp."
        ),
        70: (
            "Orchestrator tiếp nhận input gồm câu hỏi, đề bài hoặc code và chọn hướng xử lý tương ứng. Các nhánh chính gồm SUBMIT_PROBLEM, REQUEST_HINT, SUBMIT_CODE, SUBMIT_APPROACH, ASK_DIRECT_SOLUTION và ASK_THEORY."
        ),
        71: (
            "Orchestrator cũng chịu trách nhiệm kiểm soát chất lượng phản hồi. Sau khi LLM sinh phản hồi nháp, hệ thống gọi Pedagogy Critic Agent. Nếu phản hồi chưa an toàn, LLM sẽ viết lại theo `revision_instruction` trước khi gửi cho người học."
        ),
        73: (
            "State Management trong prototype được lưu trong object `LearningState` của từng agent/session. Streamlit quản lý nhiều phiên chat bằng `st.session_state`, mỗi phiên có graph và message history riêng."
        ),
        74: (
            "Trạng thái bao gồm bài toán hiện tại, dạng bài, hint level, concepts, attempts, testcase history, validation result và pedagogy flags. "
            "Các thông tin này giúp agent cá nhân hóa phản hồi theo tiến trình học."
        ),
        75: (
            "Sau mỗi lượt hội thoại, trạng thái được cập nhật ngay trong orchestrator. Ví dụ: gửi đề bài sẽ reset hint level về 0; xin gợi ý sẽ tăng hint level; gửi code sẽ cập nhật latest_validation."
        ),
        78: (
            "Module phân loại bài toán sử dụng LLM để xác định topic như dynamic_programming, graph, greedy, sliding_window hoặc sorting_searching. "
            "Kết quả được dùng để định hướng testcase và hint path."
        ),
        79: "Testcase Generator Agent",
        80: (
            "Testcase Generator Agent sinh testcase dạng JSON có cấu trúc. Testcase được chia thành các nhóm basic, edge, adversarial hoặc stress. "
            "Nếu đề bài chưa đủ rõ để xác định expected output, testcase vẫn được dùng như case tự kiểm để sinh viên trace logic."
        ),
        81: (
            "Khi sinh viên gửi code, test suite mới nhất được chuyển cho Execution Validator. Validator chỉ chấm các testcase có expected output; các testcase còn lại giúp tutor đặt câu hỏi trace tay."
        ),
        83: (
            "Hint Generation Module dùng LLM để sinh gợi ý theo trạng thái. Hint level tăng khi sinh viên xin gợi ý, nhưng policy vẫn yêu cầu không đưa full code hoặc đáp án cuối."
        ),
        84: (
            "Ở hint level thấp, hệ thống ưu tiên câu hỏi về input/output, constraint hoặc trạng thái cần lưu. Ở hint level cao hơn, gợi ý có thể cụ thể hơn nhưng vẫn yêu cầu sinh viên tự hoàn thiện lời giải."
        ),
        86: (
            "Code Analyzer tiếp nhận đoạn mã và đưa ra nhận xét về logic, độ phức tạp, edge case hoặc biến khởi tạo. Execution Validator bổ sung khả năng chạy code Python trên testcase khi có expected output."
        ),
        87: (
            "Kết quả trả về không phải bản sửa code hoàn chỉnh mà là phản hồi định hướng: case nào làm lộ lỗi, cần trace biến nào, hoặc điều kiện biên nào nên kiểm tra."
        ),
    }
    for idx, text in replacements.items():
        set_text(p[idx], text)

    # Replace the original architecture placeholder with the current architecture diagram.
    add_image_to_paragraph(p[37], ASSETS / "architecture.png", width=5.9)

    # Add current results before the original conclusion heading.
    result_heading = p[88]
    result_items = [
        "Prototype hiện tại đã chạy được bằng giao diện Streamlit. Giao diện gồm danh sách phiên chat bên trái, vùng hội thoại ở giữa, ô nhập ghim dưới cùng và panel trạng thái bên phải.",
        "Hệ thống đã tích hợp Gemini API thông qua `GeminiLlmClient`, đọc key từ file `.env` và không yêu cầu nhập key trực tiếp trên UI.",
        "LangGraph đã được tách thành nhiều node: `detect_intent`, `submit_problem`, `request_hint`, `submit_code`, `direct_solution_guard`, `submit_approach`, `ask_theory` và `finalize`. Nhờ vậy workflow agentic không còn bị gộp vào một cụm duy nhất.",
        "Quality Layer đã có testcase generator, Python execution validator và pedagogy critic. Validator hỗ trợ các trạng thái PASSED, FAILED, RUNTIME_ERROR, TIMEOUT và SKIPPED.",
        "Bộ kiểm thử tự động hiện có 14 test và đã chạy thành công với lệnh `python -m unittest discover tests`.",
    ]
    anchor = result_heading
    for text in result_items:
        anchor = insert_paragraph_after(anchor, text, style=doc.styles["Normal"])
    img_p = insert_paragraph_after(anchor, "", style=doc.styles["Normal"])
    add_image_to_paragraph(img_p, ASSETS / "demo-ui.png", width=5.9)
    cap = insert_paragraph_after(img_p, "Hình 2: Minh họa giao diện demo Streamlit của hệ thống", style=doc.styles["Normal"])
    graph_p = insert_paragraph_after(cap, "", style=doc.styles["Normal"])
    add_image_to_paragraph(graph_p, ASSETS / "langgraph.png", width=5.9)
    insert_paragraph_after(graph_p, "Hình 3: Workflow LangGraph hiện tại", style=doc.styles["Normal"])

    # Add conclusion content after the original conclusion heading.
    # The conclusion heading has shifted after inserting result content, so find it by text.
    conclusion = next(par for par in doc.paragraphs if par.text.strip() == "Kết luận")
    c1 = insert_paragraph_after(
        conclusion,
        "Đề tài đã chuyển từ ý tưởng ban đầu thành một prototype LLM-powered multi-agent tutor có khả năng chạy demo. "
        "Hệ thống sử dụng Python, Streamlit, LangGraph và Gemini API để xây dựng luồng hỗ trợ học DSA theo phương pháp Socratic.",
        style=doc.styles["Normal"],
    )
    c2 = insert_paragraph_after(
        c1,
        "So với chatbot thông thường, hệ thống có điểm mạnh ở việc quản lý trạng thái học tập và kiểm soát phản hồi bằng Quality Layer. "
        "Tutor Agent không chỉ sinh câu trả lời mà còn gọi các agent chuyên trách để sinh testcase, validate code và kiểm tra sư phạm trước khi gửi phản hồi.",
        style=doc.styles["Normal"],
    )
    c3 = insert_paragraph_after(
        c2,
        "Trong phạm vi Đồ án 2, prototype vẫn còn một số hạn chế như chưa có cơ sở dữ liệu bền vững, chưa có bộ đánh giá hội thoại lớn và sandbox mới hỗ trợ Python ở mức cơ bản. "
        "Hướng phát triển tiếp theo là lưu lịch sử học tập bằng database, nâng cấp sandbox bằng môi trường cách ly mạnh hơn và bổ sung bộ benchmark bài tập DSA để đánh giá chất lượng gợi ý.",
        style=doc.styles["Normal"],
    )
    code_p = insert_paragraph_after(c3, "", style=doc.styles["Normal"])
    add_image_to_paragraph(code_p, ASSETS / "code-pipeline.png", width=5.9)
    insert_paragraph_after(code_p, "Hình 4: Trích đoạn pipeline LLM-first trong Tutor Orchestrator", style=doc.styles["Normal"])

    doc.save(TARGET)


if __name__ == "__main__":
    main()
