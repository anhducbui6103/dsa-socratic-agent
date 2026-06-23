from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Duc_BA_20241582E.docx"
ASSETS = ROOT / "docs" / "report-assets"


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def add_para(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style=doc.styles["Normal"])


def add_heading(doc: Document, text: str, level: int) -> None:
    doc.add_paragraph(text, style=doc.styles[f"Heading {level}"])


def add_image(doc: Document, name: str, caption: str, width: float = 5.9) -> None:
    p = doc.add_paragraph(style=doc.styles["Normal"])
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(ASSETS / name), width=Inches(width))
    cap = doc.add_paragraph(caption, style=doc.styles["Normal"])
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def main() -> None:
    doc = Document(DOCX)

    # Keep the original-format front matter and design section. Rebuild the
    # corrupted/repeated implementation-result-conclusion tail from scratch.
    start = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "Triển khai hệ thống")
    for paragraph in list(doc.paragraphs[start:]):
        delete_paragraph(paragraph)

    add_heading(doc, "Triển khai hệ thống", 1)
    add_heading(doc, "Môi trường và công nghệ sử dụng", 2)
    add_para(
        doc,
        "Prototype hiện tại được triển khai bằng Python. Giao diện demo sử dụng Streamlit, workflow agentic sử dụng LangGraph, LLM provider là Gemini API và cấu hình được đọc từ file `.env` thông qua python-dotenv.",
    )
    add_para(
        doc,
        "Trong phạm vi Đồ án 2, hệ thống không xây dựng backend riêng như Node.js, NestJS hoặc FastAPI. Streamlit đảm nhiệm vai trò giao diện demo và quản lý session ở mức memory, giúp tập trung vào lõi agentic AI.",
    )
    add_para(
        doc,
        "Các thư viện chính gồm LangGraph cho workflow orchestration, Streamlit cho UI, python-dotenv cho cấu hình môi trường và unittest cho kiểm thử. API key được đặt trong `.env` dưới biến `GEMINI_API_KEY` và không nhập trực tiếp trên giao diện.",
    )

    add_heading(doc, "Triển khai AI Agent", 2)
    add_para(
        doc,
        "AI Agent được triển khai trong lớp `DsaLearningAgent`. Khi nhận input, agent gọi LLM để phân loại intent, ghi nhận input vào `student_attempts`, sau đó chọn nhánh xử lý phù hợp.",
    )
    add_para(
        doc,
        "Các hành vi chính đã được tách thành các phương thức riêng: `handle_submit_problem`, `handle_request_hint`, `handle_submit_code`, `handle_direct_solution_request`, `handle_submit_approach` và `handle_theory_question`. Việc tách này giúp LangGraph có thể gọi từng hành vi như một node riêng thay vì gom tất cả vào một node lớn.",
    )

    add_heading(doc, "Triển khai LangGraph workflow", 2)
    add_para(
        doc,
        "LangGraph không còn chỉ có một node `tutor_orchestrator`. Workflow hiện tại đã được tách thành các node: `detect_intent`, `submit_problem`, `request_hint`, `submit_code`, `direct_solution_guard`, `submit_approach`, `ask_theory` và `finalize`.",
    )
    add_para(
        doc,
        "Node `detect_intent` gọi LLM để xác định intent và cập nhật lịch sử attempt. Sau đó conditional edge định tuyến sang node tương ứng. Mỗi node xử lý một loại hành động riêng rồi chuyển về `finalize` để trả response cuối.",
    )
    add_image(doc, "langgraph.png", "Hình 2: Workflow LangGraph sau khi tách node", 5.9)

    add_heading(doc, "Triển khai State Management", 2)
    add_para(
        doc,
        "State Management trong prototype được lưu trong object `LearningState` của từng agent/session. Streamlit quản lý nhiều phiên chat bằng `st.session_state`, mỗi phiên có graph và message history riêng.",
    )
    add_para(
        doc,
        "Trạng thái bao gồm bài toán hiện tại, dạng bài, hint level, concepts, attempts, testcase history, validation result và pedagogy flags. Các thông tin này giúp agent cá nhân hóa phản hồi theo tiến trình học.",
    )

    add_heading(doc, "Triển khai các module chức năng", 2)
    add_heading(doc, "Problem Classifier", 3)
    add_para(
        doc,
        "Module phân loại bài toán sử dụng LLM để xác định topic như dynamic_programming, graph, greedy, sliding_window hoặc sorting_searching. Kết quả được dùng để định hướng testcase và hint path.",
    )
    add_heading(doc, "Testcase Generator Agent", 3)
    add_para(
        doc,
        "Testcase Generator Agent sinh testcase dạng JSON có cấu trúc. Testcase được chia thành các nhóm basic, edge, adversarial hoặc stress. Nếu đề bài chưa đủ rõ để xác định expected output, testcase vẫn được dùng như case tự kiểm để sinh viên trace logic.",
    )
    add_heading(doc, "Hint Generation Module", 3)
    add_para(
        doc,
        "Hint Generation Module dùng LLM để sinh gợi ý theo trạng thái. Hint level tăng khi sinh viên xin gợi ý, nhưng policy vẫn yêu cầu không đưa full code hoặc đáp án cuối.",
    )
    add_heading(doc, "Code Analyzer và Execution Validator", 3)
    add_para(
        doc,
        "Code Analyzer tiếp nhận đoạn mã và đưa ra nhận xét về logic, độ phức tạp, edge case hoặc biến khởi tạo. Execution Validator bổ sung khả năng chạy code Python trên testcase khi có expected output.",
    )
    add_para(
        doc,
        "Kết quả trả về không phải bản sửa code hoàn chỉnh mà là phản hồi định hướng: case nào làm lộ lỗi, cần trace biến nào, hoặc điều kiện biên nào nên kiểm tra.",
    )

    add_heading(doc, "Kết quả", 1)
    add_para(
        doc,
        "Prototype hiện tại đã chạy được bằng giao diện Streamlit. Giao diện gồm danh sách phiên chat bên trái, vùng hội thoại ở giữa, ô nhập ghim dưới cùng và panel trạng thái bên phải.",
    )
    add_para(
        doc,
        "Hệ thống đã tích hợp Gemini API thông qua `GeminiLlmClient`, đọc key từ file `.env` và không yêu cầu nhập key trực tiếp trên UI.",
    )
    add_para(
        doc,
        "Quality Layer đã có testcase generator, Python execution validator và pedagogy critic. Validator hỗ trợ các trạng thái PASSED, FAILED, RUNTIME_ERROR, TIMEOUT và SKIPPED.",
    )
    add_para(
        doc,
        "Bộ kiểm thử tự động hiện có 14 test và đã chạy thành công với lệnh `python -m unittest discover tests`.",
    )
    add_image(doc, "demo-ui.png", "Hình 3: Minh họa giao diện demo Streamlit của hệ thống", 5.9)
    add_image(doc, "code-pipeline.png", "Hình 4: Trích đoạn pipeline LLM-first trong Tutor Orchestrator", 5.9)

    add_heading(doc, "Kết luận", 1)
    add_para(
        doc,
        "Đề tài đã chuyển từ ý tưởng ban đầu thành một prototype LLM-powered multi-agent tutor có khả năng chạy demo. Hệ thống sử dụng Python, Streamlit, LangGraph và Gemini API để xây dựng luồng hỗ trợ học DSA theo phương pháp Socratic.",
    )
    add_para(
        doc,
        "So với chatbot thông thường, hệ thống có điểm mạnh ở việc quản lý trạng thái học tập và kiểm soát phản hồi bằng Quality Layer. Tutor Agent không chỉ sinh câu trả lời mà còn gọi các agent chuyên trách để sinh testcase, validate code và kiểm tra sư phạm trước khi gửi phản hồi.",
    )
    add_para(
        doc,
        "Trong phạm vi Đồ án 2, prototype vẫn còn một số hạn chế như chưa có cơ sở dữ liệu bền vững, chưa có bộ đánh giá hội thoại lớn và sandbox mới hỗ trợ Python ở mức cơ bản. Hướng phát triển tiếp theo là lưu lịch sử học tập bằng database, nâng cấp sandbox bằng môi trường cách ly mạnh hơn và bổ sung bộ benchmark bài tập DSA để đánh giá chất lượng gợi ý.",
    )

    doc.save(DOCX)


if __name__ == "__main__":
    main()
