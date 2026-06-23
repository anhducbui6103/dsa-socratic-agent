from __future__ import annotations

from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Duc_BA_20241582E_final.docx"
OUTPUT_DOCX = ROOT / "Duc_BA_20241582E_final_v2.docx"


TECH_ITEMS = [
    ("h2", "Python cho core agent"),
    (
        "p",
        "Python được sử dụng làm ngôn ngữ chính cho phần lõi của hệ thống. Các thành phần như AI Tutor Agent, LangGraph workflow, schema dữ liệu, testcase generator, validator và pedagogy review đều được tổ chức trong package Python. Việc sử dụng Python giúp hệ thống dễ tích hợp với thư viện AI, dễ viết prototype và thuận tiện cho kiểm thử tự động.",
    ),
    (
        "p",
        "Trong project, các cấu trúc dữ liệu quan trọng được định nghĩa bằng dataclass. Cách tiếp cận này giúp biểu diễn rõ các object như LearningState, AgentTurn, Classification, TestSuite, ValidationResult và PedagogyReview. Nhờ đó, output giữa các agent không chỉ là text tự do mà có cấu trúc rõ ràng để workflow có thể xử lý tiếp.",
    ),
    ("h2", "LLM provider cho năng lực hiểu ngôn ngữ"),
    (
        "p",
        "Hệ thống sử dụng LLM provider để xử lý những tác vụ cần hiểu ngôn ngữ tự nhiên và ngữ cảnh học tập. Các tác vụ này gồm nhận diện intent, phân loại dạng bài, sinh gợi ý, sinh testcase, đánh giá phản hồi sư phạm và tạo phản hồi cuối cho người học.",
    ),
    (
        "p",
        "Việc dùng LLM giúp hệ thống linh hoạt hơn so với classifier hoặc hint engine thuần rule-based. Ví dụ, cùng một bài toán quy hoạch động có thể được mô tả bằng nhiều cách khác nhau; LLM có khả năng nhận diện tín hiệu ngữ nghĩa như trạng thái, lựa chọn tối ưu, quan hệ truy hồi hoặc điều kiện biên mà không cần phụ thuộc hoàn toàn vào keyword cố định.",
    ),
    ("h2", "LangGraph cho workflow orchestration"),
    (
        "p",
        "LangGraph được sử dụng để mô hình hóa quá trình xử lý hội thoại thành một đồ thị có trạng thái. Mỗi node trong graph tương ứng với một bước hoặc một hành vi chính của agent, chẳng hạn detect_intent, submit_problem, request_hint, submit_code, direct_solution_guard, ask_theory và finalize.",
    ),
    (
        "p",
        "So với cách viết một hàm xử lý duy nhất, LangGraph giúp luồng xử lý dễ quan sát và dễ mở rộng hơn. Khi cần thêm một agent mới, hệ thống có thể bổ sung node hoặc cạnh điều kiện mà không phải viết lại toàn bộ orchestrator. Điều này đặc biệt phù hợp với kiến trúc multi-agent, nơi nhiều thành phần chuyên trách cần phối hợp theo từng ngữ cảnh.",
    ),
    (
        "p",
        "LangGraph cũng giúp báo cáo và demo rõ ràng hơn vì có thể giải thích hệ thống như một pipeline: người học gửi input, graph phát hiện intent, route sang node phù hợp, gọi các module/agent cần thiết, sau đó quay về finalize để trả phản hồi cuối.",
    ),
    ("h2", "Streamlit cho giao diện tương tác"),
    (
        "p",
        "Streamlit được sử dụng để xây dựng giao diện chat cho hệ thống. Giao diện gồm danh sách phiên chat, vùng hội thoại, ô nhập và panel trạng thái. Công nghệ này phù hợp vì cho phép xây dựng nhanh một giao diện tương tác phục vụ demo, đồng thời vẫn đủ linh hoạt để hiển thị state, testcase, validation result và Agent Trace.",
    ),
    (
        "p",
        "Trong hệ thống AI tutor, giao diện không chỉ là nơi nhập câu hỏi. UI còn giúp người học theo dõi tiến trình học, xem gợi ý theo từng lượt, quan sát các testcase gợi ý và biết hệ thống đang phân loại bài toán theo hướng nào. Vì vậy Streamlit đóng vai trò quan trọng trong việc biến core agent thành một trải nghiệm học tập có thể sử dụng được.",
    ),
    ("h2", "Quality Layer và sandbox execution"),
    (
        "p",
        "Quality Layer gồm các thành phần hỗ trợ kiểm soát chất lượng phản hồi. Testcase Generator Agent tạo testcase để người học tự kiểm hoặc để validator chạy tự động. Code Execution Validator Agent chạy code Python trên testcase có expected output và trả về trạng thái chuẩn hóa như PASSED, FAILED, RUNTIME_ERROR, TIMEOUT hoặc SKIPPED.",
    ),
    (
        "p",
        "Sandbox execution trong phiên bản hiện tại tập trung vào Python. Việc giới hạn ngôn ngữ và thời gian chạy giúp giảm rủi ro khi thực thi code do người học gửi. Kết quả chạy code không được trả thẳng dưới dạng log kỹ thuật, mà được Tutor Agent chuyển hóa thành gợi ý học tập như case nào cần trace, biến nào cần kiểm tra hoặc điều kiện biên nào có thể gây sai.",
    ),
    ("h2", "Kiểm thử và đánh giá thành phần"),
    (
        "p",
        "Hệ thống sử dụng unittest để kiểm tra các thành phần cốt lõi như intent handling, schema testcase, validation result, pedagogy review và các nhánh xử lý chính của agent. Kiểm thử giúp đảm bảo các agent trả output có cấu trúc ổn định và workflow có thể hoạt động nhất quán.",
    ),
    (
        "p",
        "Bên cạnh unit test, hệ thống còn được thiết kế để đánh giá bằng các kịch bản hội thoại. Các kịch bản này bao gồm gửi đề bài, xin gợi ý nhiều lần, gửi code sai edge case và yêu cầu lời giải trực tiếp. Đây là cách kiểm tra phù hợp với một hệ thống tutor vì chất lượng không chỉ nằm ở đúng/sai kỹ thuật mà còn ở cách phản hồi dẫn dắt người học.",
    ),
]


CONCLUSION_ITEMS = [
    ("h2", "Những kết quả đã đạt được"),
    (
        "p",
        "Hệ thống đã xây dựng được kiến trúc tổng thể cho một AI tutor hỗ trợ học Cấu trúc dữ liệu và Giải thuật theo phương pháp Socratic. Thay vì chỉ là chatbot hỏi đáp, hệ thống có AI Tutor Agent điều phối trung tâm, LangGraph workflow để định tuyến hội thoại, LearningState để lưu trạng thái học tập và Quality Layer để kiểm soát phản hồi.",
    ),
    (
        "p",
        "Một ưu điểm quan trọng của hệ thống là khả năng sử dụng LLM trong các bước cần hiểu ngữ cảnh. LLM được dùng để nhận diện intent, phân loại bài toán, sinh gợi ý, sinh testcase và đánh giá phản hồi sư phạm. Nhờ đó, hệ thống linh hoạt hơn so với cách tiếp cận rule-based thuần túy.",
    ),
    (
        "p",
        "Hệ thống cũng đã có cơ chế gợi ý nhiều cấp độ thông qua hint_level. Khi người học xin thêm gợi ý, agent không lập tức đưa lời giải mà tăng dần mức hỗ trợ. Cách làm này phù hợp với mục tiêu học tập vì người học vẫn phải tự suy luận, tự trace và tự hoàn thiện lời giải.",
    ),
    (
        "p",
        "Quality Layer là điểm mạnh khác của thiết kế. Testcase Generator Agent giúp tạo các case để kiểm tra giả định; Code Execution Validator Agent hỗ trợ phát hiện lỗi khi có expected output; Pedagogy Critic Agent kiểm tra phản hồi nháp trước khi gửi. Nhờ đó, hệ thống có thêm lớp kiểm soát thay vì phụ thuộc hoàn toàn vào một lần sinh phản hồi của LLM.",
    ),
    (
        "p",
        "Về mặt triển khai, hệ thống đã có giao diện Streamlit, workflow LangGraph tách node rõ ràng, schema dữ liệu cho các output quan trọng và các hình minh họa kiến trúc/code trong báo cáo. Đây là nền tảng đủ rõ để tiếp tục phát triển thành một trợ giảng AI hoàn chỉnh hơn.",
    ),
    ("h2", "Những hạn chế còn tồn tại"),
    (
        "p",
        "Hạn chế đầu tiên là hệ thống chưa có cơ sở dữ liệu bền vững để lưu lịch sử học tập lâu dài. LearningState hiện phù hợp cho từng phiên tương tác, nhưng để sử dụng thực tế cần bổ sung database để lưu session, tiến độ học, misconceptions và lịch sử testcase theo từng người học.",
    ),
    (
        "p",
        "Hạn chế thứ hai nằm ở sandbox execution. Phiên bản hiện tại tập trung vào Python và chỉ xử lý được các testcase có expected output rõ ràng. Với các bài toán có input/output phức tạp hoặc cần nhiều file, validator cần được mở rộng về môi trường chạy, giới hạn tài nguyên và cơ chế chuẩn hóa kết quả.",
    ),
    (
        "p",
        "Hạn chế thứ ba là testcase generator vẫn phụ thuộc vào độ rõ của đề bài và chất lượng sinh của LLM. Trong một số trường hợp, testcase chỉ nên dùng để trace tay thay vì validation tự động. Nếu ép sinh expected output khi đề chưa đủ thông tin, hệ thống có thể tạo oracle sai.",
    ),
    (
        "p",
        "Ngoài ra, hệ thống chưa có bộ benchmark lớn để đánh giá định lượng chất lượng phản hồi. Các tiêu chí như mức độ Socratic, khả năng tránh lộ lời giải, độ chính xác phân loại bài toán và hiệu quả phát hiện edge case cần được đo trên nhiều bài tập và nhiều kịch bản hội thoại hơn.",
    ),
    ("h2", "Hướng phát triển tiếp theo"),
    (
        "p",
        "Hướng phát triển đầu tiên là bổ sung database cho State & Memory Layer. Khi có lưu trữ bền vững, hệ thống có thể theo dõi tiến trình học của từng người dùng qua nhiều phiên, phát hiện misconception lặp lại và cá nhân hóa gợi ý tốt hơn.",
    ),
    (
        "p",
        "Hướng phát triển thứ hai là nâng cấp sandbox execution. Hệ thống có thể sử dụng môi trường cách ly mạnh hơn, hỗ trợ nhiều ngôn ngữ lập trình hơn, kiểm soát bộ nhớ/thời gian chạy chặt chẽ hơn và chuẩn hóa input/output theo từng dạng bài.",
    ),
    (
        "p",
        "Hướng phát triển thứ ba là xây dựng Evaluation Agent và benchmark đánh giá. Evaluation Agent có thể chấm phản hồi theo rubric như đúng hint level, không lộ full solution, có tận dụng testcase hay không và có đặt câu hỏi dẫn dắt phù hợp không.",
    ),
    (
        "p",
        "Cuối cùng, hệ thống có thể mở rộng UI để hiển thị tiến độ học tập, thống kê dạng bài đã luyện, testcase thường sai và khuyến nghị chủ đề cần ôn lại. Khi kết hợp state dài hạn, evaluation và giao diện học tập, DSA Socratic Agent có thể trở thành một trợ giảng AI cá nhân hóa cho môn Cấu trúc dữ liệu và Giải thuật.",
    ),
]


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def is_h1(paragraph) -> bool:
    return paragraph.style.name == "Heading 1"


def section_range(doc: Document, heading: str) -> tuple[int, int]:
    start = next(i for i, p in enumerate(doc.paragraphs) if is_h1(p) and p.text.strip() == heading)
    end = len(doc.paragraphs)
    for i in range(start + 1, len(doc.paragraphs)):
        if is_h1(doc.paragraphs[i]):
            end = i
            break
    return start, end


def replace_section_body(doc: Document, heading: str, items: list[tuple[str, str]]) -> None:
    start, end = section_range(doc, heading)
    heading_paragraph = doc.paragraphs[start]
    for paragraph in list(doc.paragraphs[start + 1 : end]):
        delete_paragraph(paragraph)

    for kind, text in reversed(items):
        style_name = {"h2": "Heading 2", "h3": "Heading 3", "p": "Normal"}[kind]
        paragraph = doc.add_paragraph(text, style=doc.styles[style_name])
        heading_paragraph._p.addnext(paragraph._p)


def main() -> None:
    doc = Document(DOCX)
    replace_section_body(doc, "Công nghệ sử dụng", TECH_ITEMS)
    replace_section_body(doc, "Kết luận", CONCLUSION_ITEMS)

    try:
        doc.save(DOCX)
    except PermissionError:
        doc.save(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
