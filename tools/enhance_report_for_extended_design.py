from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Duc_BA_20241582E.docx"
OUTPUT_DOCX = ROOT / "Duc_BA_20241582E_enhanced.docx"
ASSETS = ROOT / "docs" / "report-assets"


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def add_para(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style=doc.styles["Normal"])


def add_heading(doc: Document, text: str, level: int) -> None:
    doc.add_paragraph(text, style=doc.styles[f"Heading {level}"])


def add_image(doc: Document, name: str, caption: str, width: float = 5.9) -> None:
    image_path = ASSETS / name
    if not image_path.exists():
        return
    p = doc.add_paragraph(style=doc.styles["Normal"])
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(width))
    cap = doc.add_paragraph(caption, style=doc.styles["Normal"])
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def main() -> None:
    doc = Document(DOCX)

    start = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip().startswith("TÓM TẮT"))
    for paragraph in list(doc.paragraphs[start:]):
        delete_paragraph(paragraph)

    add_para(
        doc,
        "TÓM TẮT— Đề tài xây dựng một hệ thống LLM-powered Multi-Agent Tutor hỗ trợ sinh viên học môn Cấu trúc dữ liệu và Giải thuật theo phương pháp Socratic. Hệ thống không tập trung cung cấp lời giải trực tiếp, mà điều phối nhiều agent chuyên trách để phân tích đề bài, sinh gợi ý, tạo testcase, kiểm tra code và đánh giá chất lượng phản hồi trước khi gửi cho người học. Kiến trúc sử dụng Python, Streamlit, LangGraph và Gemini API, phù hợp với phạm vi Đồ án 2 nhưng vẫn thể hiện được đặc trưng của một hệ thống agentic AI có trạng thái, có workflow và có quality layer.",
    )

    add_heading(doc, "Giới thiệu", 1)
    add_para(
        doc,
        "Trong quá trình học Cấu trúc dữ liệu và Giải thuật, sinh viên thường gặp khó khăn không chỉ ở việc nhớ thuật toán, mà còn ở khả năng phân tích đề bài, nhận ra mẫu tư duy phù hợp, kiểm tra edge case và đánh giá độ phức tạp. Các công cụ hỏi đáp thông thường có thể đưa ra lời giải nhanh, nhưng điều này dễ làm người học bỏ qua quá trình hình thành tư duy giải bài.",
    )
    add_para(
        doc,
        "Đề tài DSA Socratic Agent được xây dựng nhằm giải quyết vấn đề trên bằng một hệ thống trợ giảng AI có khả năng dẫn dắt hội thoại. Thay vì trả lời như một chatbot đơn lẻ, hệ thống tổ chức quá trình hỗ trợ học tập thành workflow nhiều bước: nhận diện intent, cập nhật trạng thái học, chọn agent phù hợp, kiểm tra chất lượng và phản hồi theo phong cách Socratic.",
    )
    add_para(
        doc,
        "Điểm trọng tâm của hệ thống là kết hợp LLM với kiến trúc multi-agent. LLM được sử dụng cho những tác vụ cần hiểu ngữ cảnh như phân loại bài toán, sinh gợi ý, sinh testcase và đánh giá phản hồi sư phạm. LangGraph được dùng để biểu diễn workflow rõ ràng, giúp hệ thống dễ mở rộng, dễ debug và dễ trình bày trong báo cáo đồ án.",
    )

    add_heading(doc, "Mục tiêu", 1)
    add_heading(doc, "Mục tiêu tổng quát", 2)
    add_para(
        doc,
        "Xây dựng một prototype AI Agent hỗ trợ sinh viên học Cấu trúc dữ liệu và Giải thuật theo hướng phát triển tư duy, có khả năng duy trì trạng thái học tập, điều phối nhiều agent chuyên trách và kiểm soát phản hồi để tránh cung cấp lời giải trực tiếp trong chế độ mặc định.",
    )
    add_heading(doc, "Mục tiêu cụ thể", 2)
    add_heading(doc, "Xây dựng cơ chế điều phối bằng LangGraph", 3)
    add_para(
        doc,
        "Workflow được tách thành các node độc lập theo intent như gửi đề bài, xin gợi ý, gửi code, hỏi lý thuyết, trình bày ý tưởng hoặc yêu cầu lời giải trực tiếp. Cách tổ chức này giúp hệ thống không bị gom thành một hàm xử lý lớn, đồng thời tạo nền tảng để bổ sung agent mới trong các giai đoạn sau.",
    )
    add_heading(doc, "Sử dụng LLM cho các tác vụ cần hiểu ngữ cảnh", 3)
    add_para(
        doc,
        "Hệ thống dùng Gemini API cho intent detection, problem classification, testcase generation, final response generation và pedagogy review. Các rule/guideline nội bộ chỉ đóng vai trò hỗ trợ prompt hoặc kiểm thử, không phải nguồn sinh phản hồi chính.",
    )
    add_heading(doc, "Bổ sung Agent Trace để quan sát quá trình xử lý", 3)
    add_para(
        doc,
        "Mỗi lượt hội thoại có thể ghi lại các bước agent đã thực hiện như Intent Detector, Problem Classifier, Testcase Generator, Execution Validator và Pedagogy Critic. Agent Trace giúp minh họa rõ tính agentic của hệ thống, đồng thời hỗ trợ debug và đánh giá chất lượng trong demo.",
    )
    add_heading(doc, "Thiết kế Quality Layer cho phản hồi học tập", 3)
    add_para(
        doc,
        "Quality Layer gồm Testcase Generator Agent, Code Execution Validator Agent và Pedagogy Critic Agent. Lớp này giúp phản hồi của Tutor Agent có cơ sở kiểm chứng hơn, giảm nguy cơ lộ lời giải và định hướng người học tự kiểm tra bằng testcase hoặc trace biến.",
    )
    add_heading(doc, "Xây dựng giao diện demo phù hợp Đồ án 2", 3)
    add_para(
        doc,
        "Giao diện dùng Streamlit thay vì tách backend riêng. UI tập trung vào trải nghiệm chat, nhiều phiên học, trạng thái học tập và Agent Trace. Cách triển khai này giúp giảm độ phức tạp hạ tầng nhưng vẫn đủ để trình bày rõ kiến trúc agentic AI.",
    )

    add_heading(doc, "Thiết kế hệ thống", 1)
    add_heading(doc, "Tổng quan kiến trúc", 2)
    add_para(
        doc,
        "Hệ thống được tổ chức thành năm lớp: User Interface, Tutor Orchestration Layer, Learning Modules, Quality Layer và State & Memory Layer. AI Tutor Agent nằm ở trung tâm điều phối, còn các agent chuyên trách xử lý từng nhiệm vụ cụ thể. Mọi phản hồi gửi tới người học đều dựa trên trạng thái học tập và chính sách Socratic.",
    )
    add_image(doc, "architecture.png", "Hình 1: Tổng quan kiến trúc hệ thống multi-agent", 5.9)

    add_heading(doc, "Tutor Orchestration Layer", 2)
    add_para(
        doc,
        "Tutor Orchestration Layer chịu trách nhiệm nhận input, gọi LLM để nhận diện intent, cập nhật LearningState và định tuyến tới node xử lý phù hợp. Lớp này không sinh phản hồi trực tiếp theo kiểu một lần gọi LLM duy nhất, mà đóng vai trò điều phối chuỗi hành động của toàn bộ hệ thống.",
    )
    add_para(
        doc,
        "LangGraph biểu diễn workflow bằng các node riêng: detect_intent, submit_problem, request_hint, submit_code, direct_solution_guard, submit_approach, ask_theory và finalize. Mỗi node có trách nhiệm rõ ràng, giúp hệ thống dễ mở rộng thêm các bước như title generation, agent trace logging hoặc evaluation scoring.",
    )
    add_image(doc, "langgraph.png", "Hình 2: Workflow LangGraph tách node theo intent", 5.9)

    add_heading(doc, "Learning Modules", 2)
    add_para(
        doc,
        "Learning Modules cung cấp năng lực chuyên môn cho Tutor Agent. Problem Classifier phân loại bài toán theo chủ đề và pattern; Hint Generator sinh gợi ý theo hint_level; Code Analyzer phân tích ý tưởng, độ phức tạp và edge case; Knowledge Context cung cấp taxonomy, rubric và guideline nội bộ cho prompt.",
    )
    add_para(
        doc,
        "Trong thiết kế này, LLM là nguồn sinh phản hồi chính. Các module heuristic chỉ đóng vai trò bổ trợ: cung cấp ngữ cảnh, kiểm tra cấu trúc hoặc tạo điểm neo để unit test. Cách làm này giữ được tính linh hoạt của LLM nhưng vẫn có cấu trúc để debug.",
    )

    add_heading(doc, "Quality Layer", 2)
    add_para(
        doc,
        "Quality Layer là điểm khác biệt quan trọng so với chatbot thông thường. Testcase Generator Agent sinh các case basic, edge, adversarial hoặc stress để hỗ trợ trace. Code Execution Validator Agent chạy code Python trên testcase có expected output trong giới hạn timeout. Pedagogy Critic Agent kiểm tra phản hồi nháp trước khi gửi cho sinh viên.",
    )
    add_para(
        doc,
        "Đối với testcase, hệ thống phân biệt hai loại: testcase có expected_output dùng cho validation tự động và testcase trace-only dùng để hướng dẫn người học tự kiểm tra logic. Nếu LLM không đủ chắc chắn để sinh expected output, testcase không bị ép thành oracle sai mà được đánh dấu để dùng cho mục đích học tập.",
    )
    add_para(
        doc,
        "Pedagogy Critic đánh giá các rủi ro như lộ full code, trả lời quá trực tiếp, quá nhiều gợi ý trong một lượt hoặc không phù hợp với hint_level. Nếu phản hồi chưa an toàn, Tutor Agent gọi LLM viết lại theo revision_instruction trước khi gửi cho người học.",
    )

    add_heading(doc, "State & Memory Layer", 2)
    add_para(
        doc,
        "Mỗi phiên chat có LearningState riêng, gồm session_id, current_problem, problem_type, concepts, hint_level, student_attempts, misconceptions, generated_tests, latest_validation, pedagogy_flags và next_action. State giúp hệ thống cá nhân hóa phản hồi theo tiến trình học thay vì xử lý từng câu hỏi rời rạc.",
    )
    add_para(
        doc,
        "Bên cạnh LearningState, hệ thống còn định hướng bổ sung session title tự sinh bằng LLM. Title được tạo từ vài lượt trao đổi đầu tiên để sidebar không dùng nguyên văn message dài của sinh viên, giúp giao diện gọn hơn và phù hợp với mô hình nhiều phiên học.",
    )

    add_heading(doc, "Agent Trace và Observability", 2)
    add_para(
        doc,
        "Agent Trace ghi lại các bước xử lý chính của mỗi lượt hội thoại: node đã chạy, input/output rút gọn, trạng thái pass/fail của validator và kết quả pedagogy review. Đây không phải là chain-of-thought nội bộ của LLM, mà là trace kỹ thuật ở mức workflow để phục vụ demo, debug và đánh giá hệ thống.",
    )
    add_image(doc, "agent-trace.png", "Hình 3: Minh họa Agent Trace trong một lượt hội thoại", 5.9)

    add_heading(doc, "Triển khai hệ thống", 1)
    add_heading(doc, "Môi trường và công nghệ sử dụng", 2)
    add_para(
        doc,
        "Prototype sử dụng Python làm backend/core agent, Streamlit làm giao diện demo, LangGraph làm workflow orchestration, Gemini API làm LLM provider và python-dotenv để đọc cấu hình từ file .env. Trong phạm vi Đồ án 2, hệ thống không cần backend riêng như FastAPI hoặc Node.js vì Streamlit đã đủ cho luồng demo tương tác.",
    )
    add_para(
        doc,
        "Cấu hình LLM được đặt trong biến GEMINI_API_KEY. API key không nhập trực tiếp trên giao diện để tránh lộ khóa trong quá trình demo. Sandbox Python được xem là một thành phần luôn bật ở phía agent, còn UI chỉ hiển thị kết quả validation ở mức tóm tắt.",
    )

    add_heading(doc, "Triển khai AI Tutor Agent", 2)
    add_para(
        doc,
        "AI Tutor Agent được hiện thực bằng lớp DsaLearningAgent. Agent gồm các hành vi chính tương ứng với intent: handle_submit_problem, handle_request_hint, handle_submit_code, handle_direct_solution_request, handle_submit_approach và handle_theory_question. Việc tách method giúp LangGraph gọi từng hành vi như một node riêng.",
    )
    add_para(
        doc,
        "Khi nhận message, agent gọi LLM để detect intent, ghi nhận attempt vào state, chạy nhánh xử lý tương ứng, sinh phản hồi nháp, gọi Pedagogy Critic review và chỉ trả response cuối khi phản hồi đạt yêu cầu an toàn sư phạm.",
    )

    add_heading(doc, "Triển khai LangGraph workflow", 2)
    add_para(
        doc,
        "LangGraph định tuyến message bằng conditional edges. Node detect_intent là điểm vào của workflow; sau đó hệ thống chuyển sang node submit_problem, request_hint, submit_code, direct_solution_guard, submit_approach hoặc ask_theory. Tất cả nhánh đều hội tụ về finalize để chuẩn hóa output.",
    )
    add_para(
        doc,
        "Thiết kế này phù hợp với hướng mở rộng vì mỗi agent mới có thể được thêm thành một node hoặc một bước phụ trong node hiện có. Ví dụ, SessionTitleAgent có thể chạy sau vài lượt đầu, còn EvaluationAgent có thể chạy sau mỗi kịch bản demo để ghi điểm phản hồi.",
    )

    add_heading(doc, "Triển khai giao diện Streamlit", 2)
    add_para(
        doc,
        "Giao diện Streamlit được thiết kế theo mô hình chat: cột trái là danh sách phiên học, vùng chính là nội dung hội thoại, ô nhập được ghim dưới cùng và panel bên phải hiển thị trạng thái phiên hiện tại. Các phiên chat có title ngắn để người dùng dễ phân biệt thay vì dùng toàn bộ message đầu tiên.",
    )
    add_para(
        doc,
        "Panel bên phải tập trung vào thông tin phục vụ học tập và demo: hint level, problem type, concepts, testcase gợi ý, validation summary, pedagogy review và Agent Trace. Những thông tin kỹ thuật nhạy cảm như API key hoặc cấu hình sandbox không hiển thị như control cho người dùng cuối.",
    )
    add_image(doc, "demo-ui.png", "Hình 4: Minh họa giao diện demo Streamlit", 5.9)

    add_heading(doc, "Triển khai các agent chuyên trách", 2)
    add_heading(doc, "Intent Detector Agent", 3)
    add_para(
        doc,
        "Intent Detector Agent dùng LLM để phân loại message thành các intent chính như ASK_THEORY, SUBMIT_PROBLEM, REQUEST_HINT, SUBMIT_APPROACH, SUBMIT_CODE và ASK_DIRECT_SOLUTION. Kết quả intent quyết định nhánh xử lý trong LangGraph.",
    )
    add_heading(doc, "Problem Classifier Agent", 3)
    add_para(
        doc,
        "Problem Classifier Agent phân tích đề bài và trả về topic, pattern, confidence, key_signals và recommended_hint_path. Output có cấu trúc JSON để các module sau có thể dùng trực tiếp thay vì parse văn bản tự do.",
    )
    add_heading(doc, "Hint Generator Agent", 3)
    add_para(
        doc,
        "Hint Generator Agent sinh gợi ý theo hint_level và trạng thái học tập. Ở các mức thấp, gợi ý tập trung vào câu hỏi định hướng; ở mức cao hơn, hệ thống có thể chỉ ra cấu trúc dữ liệu hoặc invariant cần theo dõi nhưng vẫn tránh đưa full solution.",
    )
    add_heading(doc, "Testcase Generator Agent", 3)
    add_para(
        doc,
        "Testcase Generator Agent sinh TestSuite gồm nhiều GeneratedTestCase. Mỗi testcase có name, input, expected_output, category và rationale. Các testcase không chỉ dùng để chấm code mà còn là công cụ sư phạm giúp sinh viên trace tay và phát hiện edge case.",
    )
    add_heading(doc, "Code Execution Validator Agent", 3)
    add_para(
        doc,
        "Code Execution Validator Agent chạy Python code trên testcase có expected output, giới hạn timeout và chuẩn hóa kết quả thành PASSED, FAILED, RUNTIME_ERROR, TIMEOUT hoặc SKIPPED. Tutor Agent không dump log thô cho sinh viên mà chuyển kết quả thành gợi ý: case nào cần xem lại, biến nào cần trace và giả định nào có thể sai.",
    )
    add_heading(doc, "Pedagogy Critic Agent", 3)
    add_para(
        doc,
        "Pedagogy Critic Agent review phản hồi nháp theo policy Socratic. Agent này kiểm tra việc lộ lời giải, độ trực tiếp của hint, số lượng câu hỏi trong một lượt và sự phù hợp với hint_level. Đây là lớp bảo vệ cuối trước khi phản hồi được gửi ra giao diện.",
    )
    add_heading(doc, "Session Title Agent", 3)
    add_para(
        doc,
        "Session Title Agent sinh tiêu đề ngắn cho mỗi phiên chat dựa trên nội dung trao đổi ban đầu, ví dụ Maximum Subarray, DFS trên đồ thị hoặc Sliding Window. Chức năng này giúp sidebar gọn và tạo cảm giác sản phẩm hoàn chỉnh hơn trong demo.",
    )

    add_heading(doc, "Kết quả", 1)
    add_para(
        doc,
        "Hệ thống đáp ứng được các luồng demo chính của Đồ án 2: sinh viên gửi đề bài, xin gợi ý nhiều lần, gửi code để phân tích, nhận testcase tự kiểm và bị chặn khi yêu cầu lời giải trực tiếp. Luồng xử lý được thể hiện rõ bằng LangGraph và Agent Trace.",
    )
    add_para(
        doc,
        "Giao diện Streamlit hỗ trợ nhiều phiên chat, title ngắn cho từng phiên, khung nhập cố định phía dưới và panel trạng thái bên phải. Cách bố trí này giúp người dùng tập trung vào hội thoại nhưng vẫn quan sát được tiến trình học và các bước agent đã chạy.",
    )
    add_para(
        doc,
        "Quality Layer giúp hệ thống có cơ chế kiểm soát tốt hơn chatbot thông thường. Testcase Generator tạo case để trace hoặc validate; Execution Validator kiểm tra code khi đủ dữ liệu; Pedagogy Critic giảm nguy cơ phản hồi quá trực tiếp hoặc lộ full code.",
    )
    add_para(
        doc,
        "Bộ kiểm thử tự động hiện có 14 test cho các thành phần cốt lõi như intent, testcase generation, pedagogy review và execution validator. Các kịch bản tích hợp được dùng để đánh giá gồm: bài Maximum Subarray, xin hint nhiều lần, code sai edge case toàn số âm và yêu cầu full solution.",
    )
    add_image(doc, "code-pipeline.png", "Hình 5: Trích đoạn pipeline LLM-first trong Tutor Orchestrator", 5.9)

    add_heading(doc, "Đánh giá và hướng phát triển", 1)
    add_para(
        doc,
        "Trong phạm vi Đồ án 2, hệ thống ưu tiên chứng minh kiến trúc agentic AI hơn là xây dựng sản phẩm hoàn chỉnh. Việc dùng Streamlit giúp giảm chi phí triển khai backend, trong khi LangGraph và Quality Layer vẫn cho thấy rõ quá trình điều phối nhiều agent.",
    )
    add_para(
        doc,
        "Hạn chế chính nằm ở sandbox mới hỗ trợ Python cơ bản, testcase expected output phụ thuộc vào độ rõ của đề bài và state hiện được lưu ở memory thay vì database bền vững. Ngoài ra, cần bổ sung bộ benchmark lớn hơn để đo chất lượng gợi ý, khả năng tránh lộ lời giải và độ chính xác của classifier.",
    )
    add_para(
        doc,
        "Hướng phát triển tiếp theo gồm lưu session vào database, mở rộng sandbox bằng môi trường cách ly mạnh hơn, bổ sung evaluation dashboard, hỗ trợ thêm ngôn ngữ lập trình và xây dựng bộ dữ liệu bài tập DSA để đánh giá định lượng chất lượng tutor.",
    )

    add_heading(doc, "Kết luận", 1)
    add_para(
        doc,
        "DSA Socratic Agent là một prototype multi-agent tutor cho môn Cấu trúc dữ liệu và Giải thuật. Hệ thống kết hợp LLM, LangGraph, Streamlit và Quality Layer để tạo ra trải nghiệm học tập định hướng tư duy thay vì cung cấp đáp án trực tiếp.",
    )
    add_para(
        doc,
        "Điểm mạnh của thiết kế là có workflow rõ ràng, có state học tập, có các agent chuyên trách và có lớp kiểm soát sư phạm trước khi phản hồi. Nhờ đó, hệ thống phù hợp để trình bày trong Đồ án 2 như một minh chứng cho cách ứng dụng LLM vào bài toán giáo dục có kiểm soát.",
    )

    try:
        doc.save(DOCX)
    except PermissionError:
        doc.save(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
